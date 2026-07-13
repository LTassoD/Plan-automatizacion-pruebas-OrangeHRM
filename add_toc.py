from docx import Document
from docx.shared import Pt
from lxml import etree
import os

NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

def make_toc_paragraph(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    runs_data = [
        ('fldChar', 'begin', None),
        ('instrText', None, r'TOC \o "1-2" \h \z \u'),
        ('fldChar', 'separate', None),
        ('text', None, '[Actualizar indice: clic derecho > Actualizar campo]'),
        ('fldChar', 'end', None),
    ]

    for kind, fld_type, txt in runs_data:
        run = p.add_run()
        run_element = run._r
        rpr = etree.SubElement(run_element, f'{NS}rPr')
        if kind == 'fldChar':
            fld = etree.SubElement(run_element, f'{NS}fldChar')
            fld.set(f'{NS}fldCharType', fld_type)
        elif kind == 'instrText':
            instr = etree.SubElement(run_element, f'{NS}instrText')
            instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            instr.text = txt
        elif kind == 'text':
            t = etree.SubElement(run_element, f'{NS}t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = txt
    return p

def add_toc_to_doc(path):
    doc = Document(path)
    index_para = None
    index_para_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.upper().replace('\u00cd', 'I').startswith('INDICE') and p.style.name.startswith('Heading'):
            index_para = p
            index_para_idx = i
            break

    print(f'{os.path.basename(path)}: INDICE at para {index_para_idx}')

    if index_para:
        paras_to_remove = []
        for i in range(index_para_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            text = p.text.strip()
            if not text:
                continue
            if p.style.name.startswith('Heading') and 'ndice' not in text.upper().replace('\u00cd', 'I'):
                break
            paras_to_remove.append(p)
        print(f'  Removing {len(paras_to_remove)} text TOC paragraphs')
        for p in paras_to_remove:
            p._element.getparent().remove(p._element)

        toc_p = make_toc_paragraph(doc)
        toc_element = toc_p._element
        doc.element.body.remove(toc_element)
        index_para._element.addnext(toc_element)
        doc.save(path)
        print(f'  Saved')
    else:
        print(f'  No INDICE heading found (adding new section)')
        # Add INDICE after last paragraph before first Heading 1
        first_h1_idx = -1
        for i, p in enumerate(doc.paragraphs):
            if p.style.name == 'Heading 1' and p.text.strip():
                first_h1_idx = i
                break
        if first_h1_idx > 0:
            # Create INDICE heading
            heading_p = doc.add_paragraph()
            heading_p.style = doc.styles['Heading 1']
            run = heading_p.add_run('ÍNDICE')
            run.bold = True
            run.font.size = Pt(14)
            heading_element = heading_p._element
            doc.element.body.remove(heading_element)

            toc_p = make_toc_paragraph(doc)
            toc_element = toc_p._element
            doc.element.body.remove(toc_element)

            first_h1 = doc.paragraphs[first_h1_idx]._element
            first_h1.addprevious(heading_element)
            heading_element.addnext(toc_element)
            doc.save(path)
            print(f'  Created INDICE + TOC before para {first_h1_idx}')
        else:
            print(f'  No Heading 1 found to anchor INDICE')

base = r'C:\Users\User\Desktop\Duoc\Semestre V\Automatizacion de Pruebas\AutomatizacionPruebasPython'
add_toc_to_doc(os.path.join(base, 'INFORME_FINAL_EFT.docx'))
add_toc_to_doc(os.path.join(base, 'ET_FORMATO_INFORME_API0101.docx'))
